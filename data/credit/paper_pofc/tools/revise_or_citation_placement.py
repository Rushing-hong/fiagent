from __future__ import annotations

import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(r"D:\cvnotenough\data\experiments\a1_oracle\paper\OR_submission_expanded")
SOURCE = ROOT / "sharp_scale_frontiers_OR_main_regular.docx"


def clone_run_format(source_run, target_run) -> None:
    source_rpr = source_run._element.rPr
    if source_rpr is None:
        return
    target_rpr = target_run._element.rPr
    if target_rpr is not None:
        target_run._element.remove(target_rpr)
    target_run._element.insert(0, deepcopy(source_rpr))


def append_sentence(paragraph, sentence: str) -> bool:
    if sentence in paragraph.text:
        return False
    run = paragraph.add_run(" " + sentence)
    if len(paragraph.runs) >= 2:
        clone_run_format(paragraph.runs[-2], run)
    return True


def prepend_sentence(paragraph, sentence: str) -> bool:
    if sentence in paragraph.text:
        return False
    run_element = OxmlElement("w:r")
    paragraph._p.insert(0, run_element)
    from docx.text.run import Run

    run = Run(run_element, paragraph)
    run.text = sentence + " "
    if len(paragraph.runs) >= 2:
        clone_run_format(paragraph.runs[1], run)
    return True


def find_unique(doc: Document, startswith: str, contains: str | None = None):
    matches = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith(startswith)
        and (contains is None or contains in paragraph.text)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph starting with {startswith!r}; found {len(matches)}"
        )
    return matches[0]


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = ROOT / f"sharp_scale_frontiers_OR_main_regular_before_citation_redistribution_{stamp}.docx"
    shutil.copy2(SOURCE, backup)

    doc = Document(SOURCE)
    changes = 0

    changes += append_sentence(
        find_unique(doc, "We ask when descending scale order is optimal"),
        "The distinction follows the classical order-selection and descending-optimality line (Hill 1983; Hill and Hordijk 1985), while instance-wise algorithms instead optimize a supplied heterogeneous list (Agrawal, Sethuraman, and Zhang 2020; Liu et al. 2021).",
    )

    changes += append_sentence(
        find_unique(doc, "The paper contributes to exact order selection"),
        "The information structure also separates the model from costly-inspection search, where inspection costs and recall determine index or reservation rules (Weitzman 1979; Olszewski and Weber 2015).",
    )

    changes += prepend_sentence(
        find_unique(doc, "", "Hill (1983) showed that adaptive choice"),
        "The recursion is the standard finite-horizon stopping-value formulation of Snell (1952) and Chow, Robbins, and Siegmund (1971).",
    )

    descending_definition = find_unique(
        doc, "Following Hill and Hordijk (1985),", "is descending-optimal if"
    )
    changes += prepend_sentence(
        descending_definition,
        "Following Hill and Hordijk (1985),",
    )
    for run in descending_definition.runs:
        if run.text.startswith("The family"):
            run.text = "the family" + run.text[len("The family") :]
            changes += 1
            break

    changes += append_sentence(
        find_unique(doc, "A -point distribution has at most"),
        "This affine-composition viewpoint is related to ordering monotone piecewise-linear maps (Kawase, Makino, and Seimi 2018), while common-base scale conjugacy imposes the linked coefficients exploited below.",
    )

    changes += append_sentence(
        find_unique(doc, "The compression is also algorithmic."),
        "The comparison is with instance-wise ordering methods, which are exact for important two-point classes and approximate or harder in richer heterogeneous settings (Agrawal, Sethuraman, and Zhang 2020; Fu, Li, and Xu 2018; Liu et al. 2021).",
    )

    changes += append_sentence(
        find_unique(doc, "For a positive scale menu"),
        "Wasserstein sensitivity and distributionally robust guarantees are established tools for stochastic optimization (Mohajerin Esfahani and Kuhn 2018; Bartl and Wiesel 2023; Gao 2023); the object here is the gap between observation orders after optimization over permutations and normalized menus.",
    )

    temp = SOURCE.with_suffix(".citation_revised.tmp.docx")
    doc.save(temp)
    temp.replace(SOURCE)
    print(f"Updated {SOURCE}")
    print(f"Backup {backup}")
    print(f"Paragraph edits applied: {changes}")


if __name__ == "__main__":
    main()
